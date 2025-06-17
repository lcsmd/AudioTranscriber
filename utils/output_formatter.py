import os
import tempfile
import logging
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from docx import Document
from docx.shared import Inches
import markdown
import re

logger = logging.getLogger(__name__)

def format_as_text(content, metadata=None):
    """
    Format content as plain text
    
    Args:
        content (str): Content to format
        metadata (dict): Optional metadata to include
    
    Returns:
        str: Formatted text content
    """
    output = ""
    
    if metadata:
        output += "=== TRANSCRIPTION DETAILS ===\n"
        if metadata.get('filename'):
            output += f"Source File: {metadata['filename']}\n"
        if metadata.get('processing_time'):
            output += f"Processing Time: {metadata['processing_time']/1000:.1f} seconds\n"
        if metadata.get('language'):
            output += f"Language: {metadata['language']}\n"
        if metadata.get('created_at'):
            output += f"Processed: {metadata['created_at']}\n"
        output += "=" * 30 + "\n\n"
    
    output += content
    
    return output

def format_as_markdown(content, metadata=None):
    """
    Format content as Markdown
    
    Args:
        content (str): Content to format
        metadata (dict): Optional metadata to include
    
    Returns:
        str: Formatted Markdown content
    """
    output = ""
    
    if metadata:
        output += "# Transcription Details\n\n"
        if metadata.get('filename'):
            output += f"**Source File:** {metadata['filename']}\n\n"
        if metadata.get('processing_time'):
            output += f"**Processing Time:** {metadata['processing_time']/1000:.1f} seconds\n\n"
        if metadata.get('language'):
            output += f"**Language:** {metadata['language']}\n\n"
        if metadata.get('created_at'):
            output += f"**Processed:** {metadata['created_at']}\n\n"
        output += "---\n\n"
    
    output += "## Transcription\n\n"
    
    # Convert basic formatting if present
    formatted_content = content
    
    # Convert line breaks to proper markdown paragraphs
    paragraphs = formatted_content.split('\n\n')
    formatted_paragraphs = []
    
    for para in paragraphs:
        if para.strip():
            # Clean up extra whitespace
            clean_para = ' '.join(para.split())
            formatted_paragraphs.append(clean_para)
    
    output += '\n\n'.join(formatted_paragraphs)
    
    return output

def create_word_document(content, metadata=None, output_path=None):
    """
    Create a Word document from content
    
    Args:
        content (str): Content to include
        metadata (dict): Optional metadata
        output_path (str): Path to save the document
    
    Returns:
        str: Path to the created document
    """
    if output_path is None:
        import uuid
        filename = f"transcription_{uuid.uuid4().hex[:8]}.docx"
        output_path = os.path.join(tempfile.gettempdir(), filename)
    
    try:
        logger.debug(f"Creating Word document: {output_path}")
        
        doc = Document()
        
        # Add title
        title = doc.add_heading('Audio Transcription', 0)
        
        # Add metadata if provided
        if metadata:
            doc.add_heading('Document Details', level=1)
            
            if metadata.get('filename'):
                p = doc.add_paragraph()
                p.add_run('Source File: ').bold = True
                p.add_run(metadata['filename'])
            
            if metadata.get('processing_time'):
                p = doc.add_paragraph()
                p.add_run('Processing Time: ').bold = True
                p.add_run(f"{metadata['processing_time']/1000:.1f} seconds")
            
            if metadata.get('language'):
                p = doc.add_paragraph()
                p.add_run('Language: ').bold = True
                p.add_run(metadata['language'])
            
            if metadata.get('created_at'):
                p = doc.add_paragraph()
                p.add_run('Processed: ').bold = True
                p.add_run(str(metadata['created_at']))
        
        # Add content
        doc.add_heading('Transcription', level=1)
        
        # Split content into paragraphs and add them
        paragraphs = content.split('\n\n')
        for para in paragraphs:
            if para.strip():
                doc.add_paragraph(para.strip())
        
        # Save the document
        doc.save(output_path)
        
        logger.debug(f"Successfully created Word document: {output_path}")
        return output_path
        
    except Exception as e:
        logger.error(f"Error creating Word document: {str(e)}")
        raise Exception(f"Failed to create Word document: {str(e)}")

def create_pdf_document(content, metadata=None, output_path=None):
    """
    Create a PDF document from content
    
    Args:
        content (str): Content to include
        metadata (dict): Optional metadata
        output_path (str): Path to save the document
    
    Returns:
        str: Path to the created document
    """
    if output_path is None:
        import uuid
        filename = f"transcription_{uuid.uuid4().hex[:8]}.pdf"
        output_path = os.path.join(tempfile.gettempdir(), filename)
    
    try:
        logger.debug(f"Creating PDF document: {output_path}")
        
        doc = SimpleDocTemplate(output_path, pagesize=letter)
        styles = getSampleStyleSheet()
        
        # Create custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=30,
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            spaceAfter=12,
        )
        
        story = []
        
        # Add title
        story.append(Paragraph("Audio Transcription", title_style))
        story.append(Spacer(1, 20))
        
        # Add metadata if provided
        if metadata:
            story.append(Paragraph("Document Details", heading_style))
            
            if metadata.get('filename'):
                story.append(Paragraph(f"<b>Source File:</b> {metadata['filename']}", styles['Normal']))
            
            if metadata.get('processing_time'):
                story.append(Paragraph(f"<b>Processing Time:</b> {metadata['processing_time']/1000:.1f} seconds", styles['Normal']))
            
            if metadata.get('language'):
                story.append(Paragraph(f"<b>Language:</b> {metadata['language']}", styles['Normal']))
            
            if metadata.get('created_at'):
                story.append(Paragraph(f"<b>Processed:</b> {metadata['created_at']}", styles['Normal']))
            
            story.append(Spacer(1, 20))
        
        # Add content
        story.append(Paragraph("Transcription", heading_style))
        
        # Split content into paragraphs and add them
        paragraphs = content.split('\n\n')
        for para in paragraphs:
            if para.strip():
                # Escape HTML characters for ReportLab
                clean_para = para.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                story.append(Paragraph(clean_para, styles['Normal']))
                story.append(Spacer(1, 12))
        
        # Build the PDF
        doc.build(story)
        
        logger.debug(f"Successfully created PDF document: {output_path}")
        return output_path
        
    except Exception as e:
        logger.error(f"Error creating PDF document: {str(e)}")
        raise Exception(f"Failed to create PDF document: {str(e)}")

def generate_output_file(content, format_type, metadata=None, output_dir=None):
    """
    Generate an output file in the specified format
    
    Args:
        content (str): Content to format
        format_type (str): Output format (text, markdown, word, pdf)
        metadata (dict): Optional metadata
        output_dir (str): Directory to save the file
    
    Returns:
        str: Path to the generated file
    """
    if output_dir is None:
        output_dir = tempfile.gettempdir()
    
    import uuid
    base_filename = f"transcription_{uuid.uuid4().hex[:8]}"
    
    format_type = format_type.lower()
    
    if format_type == 'text':
        filename = f"{base_filename}.txt"
        output_path = os.path.join(output_dir, filename)
        formatted_content = format_as_text(content, metadata)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(formatted_content)
        
        return output_path
        
    elif format_type == 'markdown':
        filename = f"{base_filename}.md"
        output_path = os.path.join(output_dir, filename)
        formatted_content = format_as_markdown(content, metadata)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(formatted_content)
        
        return output_path
        
    elif format_type == 'word':
        filename = f"{base_filename}.docx"
        output_path = os.path.join(output_dir, filename)
        return create_word_document(content, metadata, output_path)
        
    elif format_type == 'pdf':
        filename = f"{base_filename}.pdf"
        output_path = os.path.join(output_dir, filename)
        return create_pdf_document(content, metadata, output_path)
        
    else:
        raise ValueError(f"Unsupported output format: {format_type}")

def get_supported_formats():
    """
    Get list of supported output formats
    
    Returns:
        dict: Dictionary of supported formats with descriptions
    """
    return {
        'text': {
            'name': 'Plain Text',
            'extension': '.txt',
            'description': 'Simple text format with basic formatting'
        },
        'markdown': {
            'name': 'Markdown',
            'extension': '.md',
            'description': 'Markdown format for easy web publishing'
        },
        'word': {
            'name': 'Microsoft Word',
            'extension': '.docx',
            'description': 'Microsoft Word document format'
        },
        'pdf': {
            'name': 'PDF Document',
            'extension': '.pdf',
            'description': 'Portable Document Format for universal viewing'
        }
    }