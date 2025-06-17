import os
import logging
import tempfile
from PyPDF2 import PdfReader
from docx import Document
import re

logger = logging.getLogger(__name__)

def extract_text_from_pdf(pdf_path):
    """
    Extract text content from a PDF file
    
    Args:
        pdf_path (str): Path to the PDF file
    
    Returns:
        str: Extracted text content
    """
    try:
        logger.debug(f"Extracting text from PDF: {pdf_path}")
        
        with open(pdf_path, 'rb') as file:
            pdf_reader = PdfReader(file)
            text_content = ""
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    text_content += f"\n--- Page {page_num + 1} ---\n"
                    text_content += page_text
                except Exception as e:
                    logger.warning(f"Failed to extract text from page {page_num + 1}: {str(e)}")
            
            # Clean up the text
            text_content = clean_extracted_text(text_content)
            
            logger.debug(f"Successfully extracted {len(text_content)} characters from PDF")
            return text_content
            
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {str(e)}")
        raise Exception(f"Failed to extract text from PDF: {str(e)}")

def extract_text_from_docx(docx_path):
    """
    Extract text content from a DOCX file
    
    Args:
        docx_path (str): Path to the DOCX file
    
    Returns:
        str: Extracted text content
    """
    try:
        logger.debug(f"Extracting text from DOCX: {docx_path}")
        
        doc = Document(docx_path)
        text_content = ""
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            text_content += para.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text_content += cell.text + "\t"
                text_content += "\n"
        
        # Clean up the text
        text_content = clean_extracted_text(text_content)
        
        logger.debug(f"Successfully extracted {len(text_content)} characters from DOCX")
        return text_content
        
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {str(e)}")
        raise Exception(f"Failed to extract text from DOCX: {str(e)}")

def extract_text_from_txt(txt_path):
    """
    Extract text content from a TXT file
    
    Args:
        txt_path (str): Path to the TXT file
    
    Returns:
        str: Extracted text content
    """
    try:
        logger.debug(f"Reading text from TXT: {txt_path}")
        
        # Try different encodings
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                with open(txt_path, 'r', encoding=encoding) as file:
                    text_content = file.read()
                    logger.debug(f"Successfully read TXT file with {encoding} encoding")
                    return clean_extracted_text(text_content)
            except UnicodeDecodeError:
                continue
        
        # If all encodings fail, try reading as binary and decoding with errors='ignore'
        with open(txt_path, 'rb') as file:
            text_content = file.read().decode('utf-8', errors='ignore')
            logger.warning("Used fallback encoding with error handling")
            return clean_extracted_text(text_content)
        
    except Exception as e:
        logger.error(f"Error reading text file: {str(e)}")
        raise Exception(f"Failed to read text file: {str(e)}")

def clean_extracted_text(text):
    """
    Clean and normalize extracted text
    
    Args:
        text (str): Raw extracted text
    
    Returns:
        str: Cleaned text
    """
    # Remove excessive whitespace
    text = re.sub(r'\n\s*\n\s*\n', '\n\n', text)  # Replace multiple newlines with double newlines
    text = re.sub(r'[ \t]+', ' ', text)  # Replace multiple spaces/tabs with single space
    text = text.strip()
    
    return text

def process_document(file_path, file_type):
    """
    Process a document file and extract its text content
    
    Args:
        file_path (str): Path to the document file
        file_type (str): Type of the document (pdf, docx, txt)
    
    Returns:
        str: Extracted text content
    """
    file_type = file_type.lower()
    
    if file_type == 'pdf':
        return extract_text_from_pdf(file_path)
    elif file_type == 'docx':
        return extract_text_from_docx(file_path)
    elif file_type == 'txt':
        return extract_text_from_txt(file_path)
    else:
        raise ValueError(f"Unsupported document type: {file_type}")

def get_document_info(file_path, file_type):
    """
    Get basic information about a document
    
    Args:
        file_path (str): Path to the document file
        file_type (str): Type of the document
    
    Returns:
        dict: Document information
    """
    try:
        file_size = os.path.getsize(file_path)
        file_name = os.path.basename(file_path)
        
        info = {
            'filename': file_name,
            'file_type': file_type,
            'file_size': file_size,
            'page_count': None,
            'word_count': None,
            'character_count': None,
        }
        
        # Get additional info based on file type
        if file_type.lower() == 'pdf':
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PdfReader(file)
                    info['page_count'] = len(pdf_reader.pages)
            except:
                pass
                
        elif file_type.lower() == 'docx':
            try:
                doc = Document(file_path)
                info['page_count'] = len(doc.paragraphs)  # Approximate
            except:
                pass
        
        # Get text stats
        try:
            text_content = process_document(file_path, file_type)
            info['character_count'] = len(text_content)
            info['word_count'] = len(text_content.split())
        except:
            pass
            
        return info
        
    except Exception as e:
        logger.error(f"Error getting document info: {str(e)}")
        return {
            'filename': os.path.basename(file_path),
            'file_type': file_type,
            'error': str(e)
        }